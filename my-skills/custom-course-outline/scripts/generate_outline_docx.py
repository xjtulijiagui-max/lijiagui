#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
课程大纲 Word 文档生成器
将 JSON 格式的课程大纲数据渲染为标准格式的 .docx 文件。

用法:
    python generate_outline_docx.py --input data.json --output 大纲.docx
"""

import argparse
import json
import sys
import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── 字体常量 ──
FONT_CN = "微软雅黑"       # 中文字体
FONT_CN_BODY = "宋体"      # 正文中文字体
FONT_EN = "Calibri"        # 西文字体


def set_cell_background(cell, color_hex):
    """给表格单元格设置背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """设置单元格边框。kwargs: top, bottom, left, right -> dict(sz, val, color)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", "4")}" w:space="0" '
            f'w:color="{attrs.get("color", "999999")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_run_font(run, size=None, bold=None, color=None, font_cn=FONT_CN_BODY, font_en=FONT_EN):
    """统一设置 run 的字体（中西文分别设置）"""
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run.font.name = font_en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_cn)
    rFonts.set(qn("w:ascii"), font_en)
    rFonts.set(qn("w:hAnsi"), font_en)


def add_styled_paragraph(doc, text, size=11, bold=False, color=None,
                         alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         space_before=0, space_after=4, font_cn=FONT_CN_BODY):
    """添加一个带样式的段落"""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, font_cn=font_cn)
    return p


def add_section_heading(doc, text):
    """添加板块标题（如「学习目标」「课程特色」）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=14, bold=True, color=RGBColor(0x1a, 0x1a, 0x2e), font_cn=FONT_CN)
    return p


def add_bullet(doc, text, size=10.5):
    """添加 bullet 段落（以 · 开头）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("· " + text)
    set_run_font(run, size=size, font_cn=FONT_CN_BODY)
    return p


def add_multiline_cell(cell, content, bold_name=False, name_text=None):
    """
    往单元格写入多行内容。
    content 中的 \n 分行，每行根据前缀标记设置样式：
      ▶  -> 主要点（正常字号）
      —  -> 子要点（稍小，缩进）
      其他 -> 普通文本
    name_text 如果提供，会作为单元格开头的加粗主题名。
    """
    cell.text = ""  # 清空默认空段落
    lines = content.split("\n")

    first_para = True
    if name_text:
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(name_text)
        set_run_font(run, size=10.5, bold=True, font_cn=FONT_CN)
        first_para = False

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if first_para and not name_text:
            p = cell.paragraphs[0]
            first_para = False
        else:
            p = cell.add_paragraph()

        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.3

        stripped = line.lstrip()
        if stripped.startswith("▶"):
            p.paragraph_format.left_indent = Cm(0.3)
            run = p.add_run(stripped)
            set_run_font(run, size=10, font_cn=FONT_CN_BODY)
        elif stripped.startswith("—"):
            p.paragraph_format.left_indent = Cm(0.8)
            run = p.add_run(stripped)
            set_run_font(run, size=9.5, color=RGBColor(0x55, 0x55, 0x55), font_cn=FONT_CN_BODY)
        elif stripped.startswith("📺") or stripped.startswith("🔧"):
            run = p.add_run(stripped)
            set_run_font(run, size=10, bold=True, font_cn=FONT_CN)
        else:
            run = p.add_run(stripped)
            set_run_font(run, size=10, font_cn=FONT_CN_BODY)


def build_basic_info_table(doc, info):
    """构建「课程基本信息」表格"""
    fields = [
        ("【课程名称】", info.get("course_name", "")),
        ("【课程时长】", info.get("duration", "")),
        ("【课程对象】", info.get("audience", "")),
        ("【授课形式】", info.get("format", "")),
        ("【课程主线】", info.get("main_line", "")),
        ("【课程说明】", info.get("description", "")),
    ]

    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # 设置列宽
    widths = [Cm(3), Cm(5), Cm(3), Cm(5)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

    for idx, (label, value) in enumerate(fields):
        row_idx = idx // 2
        col_group = idx % 2  # 0 -> cols 0,1;  1 -> cols 2,3
        label_col = col_group * 2
        value_col = label_col + 1

        label_cell = table.cell(row_idx, label_col)
        value_cell = table.cell(row_idx, value_col)

        # 标签格
        label_cell.text = ""
        lp = label_cell.paragraphs[0]
        lp.paragraph_format.space_before = Pt(2)
        lp.paragraph_format.space_after = Pt(2)
        lr = lp.add_run(label)
        set_run_font(lr, size=10, bold=True, color=RGBColor(0x33, 0x33, 0x33), font_cn=FONT_CN)
        set_cell_background(label_cell, "F0F0F5")

        # 值格
        value_cell.text = ""
        vp = value_cell.paragraphs[0]
        vp.paragraph_format.space_before = Pt(2)
        vp.paragraph_format.space_after = Pt(2)
        vr = vp.add_run(value)
        set_run_font(vr, size=10, font_cn=FONT_CN_BODY)

    # 给所有单元格加边框
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={"sz": "4", "color": "CCCCCC"},
                bottom={"sz": "4", "color": "CCCCCC"},
                left={"sz": "4", "color": "CCCCCC"},
                right={"sz": "4", "color": "CCCCCC"})

    return table


def build_module_table(doc, module):
    """构建单个模块的表格"""
    topics = module.get("topics", [])
    num_rows = 1 + len(topics)  # 1 header + topics

    table = doc.add_table(rows=num_rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # 列宽
    for row in table.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(11.5)

    # 第1行：模块标题（合并单元格）
    header_cell = table.cell(0, 0)
    header_cell.merge(table.cell(0, 1))
    header_cell.text = ""
    hp = header_cell.paragraphs[0]
    hp.paragraph_format.space_before = Pt(4)
    hp.paragraph_format.space_after = Pt(4)
    hr = hp.add_run(f"{module['header']}  {module['time']}")
    set_run_font(hr, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), font_cn=FONT_CN)
    # 根据类型选背景色
    mtype = module.get("type", "module")
    bg = "4A4A6A" if mtype == "module" else ("5B7C99" if mtype == "opening" else "6B5B7C")
    set_cell_background(header_cell, bg)

    # 后续行：主题
    for i, topic in enumerate(topics):
        row_idx = i + 1
        name_cell = table.cell(row_idx, 0)
        content_cell = table.cell(row_idx, 1)

        # 主题名
        name_cell.text = ""
        np = name_cell.paragraphs[0]
        np.paragraph_format.space_before = Pt(3)
        np.paragraph_format.space_after = Pt(3)
        nr = np.add_run(topic.get("name", ""))
        set_run_font(nr, size=10, bold=True, color=RGBColor(0x33, 0x33, 0x33), font_cn=FONT_CN)
        set_cell_background(name_cell, "F5F5FA")

        # 内容
        add_multiline_cell(content_cell, topic.get("content", ""))

    # 所有单元格边框
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={"sz": "4", "color": "D0D0D0"},
                bottom={"sz": "4", "color": "D0D0D0"},
                left={"sz": "4", "color": "D0D0D0"},
                right={"sz": "4", "color": "D0D0D0"})

    # 模块间空行
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def build_versions_table(doc, versions):
    """构建「版本组合说明」表格"""
    if not versions:
        return

    table = doc.add_table(rows=len(versions), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for row in table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(9)
        row.cells[2].width = Cm(3)

    for i, ver in enumerate(versions):
        name_cell = table.cell(i, 0)
        desc_cell = table.cell(i, 1)
        dur_cell = table.cell(i, 2)

        for cell, key, bold in [
            (name_cell, "name", True),
            (desc_cell, "description", False),
            (dur_cell, "duration", False),
        ]:
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(ver.get(key, ""))
            set_run_font(run, size=10, bold=bold, font_cn=FONT_CN_BODY)

    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={"sz": "4", "color": "CCCCCC"},
                bottom={"sz": "4", "color": "CCCCCC"},
                left={"sz": "4", "color": "CCCCCC"},
                right={"sz": "4", "color": "CCCCCC"})


def generate_docx(data, output_path):
    """主生成函数"""
    doc = Document()

    # ── 页面设置 ──
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── 设置默认样式 ──
    style = doc.styles["Normal"]
    style.font.name = FONT_EN
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN_BODY)

    # ════════ ① 标题 + 副标题 ════════
    add_styled_paragraph(doc, data.get("title", ""),
                         size=22, bold=True,
                         color=RGBColor(0x1a, 0x1a, 0x2e),
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=12, space_after=6,
                         font_cn=FONT_CN)

    tagline = data.get("tagline", "")
    if tagline:
        add_styled_paragraph(doc, tagline,
                             size=11, bold=False,
                             color=RGBColor(0x66, 0x66, 0x66),
                             alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_after=12, font_cn=FONT_CN)

    version_label = data.get("version_label", "")
    if version_label:
        add_styled_paragraph(doc, version_label,
                             size=14, bold=True,
                             color=RGBColor(0x33, 0x33, 0x33),
                             alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_after=16, font_cn=FONT_CN)

    # ════════ ② 课程基本信息 ════════
    add_section_heading(doc, "课程基本信息")
    build_basic_info_table(doc, data.get("basic_info", {}))

    # ════════ ③ 学习目标 ════════
    objectives = data.get("learning_objectives", [])
    if objectives:
        add_section_heading(doc, "学习目标")
        for obj in objectives:
            add_bullet(doc, obj)

    # ════════ ④ 课程模块结构 ════════
    modules = data.get("modules", [])
    if modules:
        add_section_heading(doc, "课程模块结构")
        for module in modules:
            build_module_table(doc, module)

    # ════════ ⑤ 讲师演示说明 ════════
    demo_notes = data.get("demo_notes", [])
    if demo_notes:
        add_section_heading(doc, "讲师演示说明")
        for note in demo_notes:
            add_bullet(doc, note)

    # ════════ ⑥ 课程特色 ════════
    features = data.get("features", [])
    if features:
        add_section_heading(doc, "课程特色")
        for feat in features:
            add_bullet(doc, feat)

    # ════════ ⑦ 版本组合说明 ════════
    versions = data.get("versions", [])
    if versions:
        add_section_heading(doc, "版本组合说明")
        build_versions_table(doc, versions)
        # 备注
        add_styled_paragraph(doc,
                             "*备注：如需其他时间版本，可根据具体内容另行适配输出。*",
                             size=9, color=RGBColor(0x99, 0x99, 0x99),
                             space_before=8, font_cn=FONT_CN_BODY)

    # ── 保存 ──
    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description="课程大纲 Word 生成器")
    parser.add_argument("--input", required=True, help="JSON 数据文件路径")
    parser.add_argument("--output", required=True, help="输出 .docx 路径")
    args = parser.parse_args()

    with open(args.input, "r", encoding="utf-8") as f:
        data = json.load(f)

    result = generate_docx(data, args.output)
    print(f"OK: {result}")


if __name__ == "__main__":
    main()
